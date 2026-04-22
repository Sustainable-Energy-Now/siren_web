"""
Generate a demonstration script Word document for the Transmission module.

Usage:
    python docs/generate_transmission_demo.py

Output:
    docs/transmission_demo_script.docx
"""

import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------
COLOUR_HEADING    = RGBColor(0x1A, 0x53, 0x76)   # dark blue
COLOUR_SUBHEADING = RGBColor(0x2E, 0x86, 0xAB)   # mid blue
COLOUR_ACCENT     = RGBColor(0x27, 0xAE, 0x60)   # green
COLOUR_WARNING    = RGBColor(0xE6, 0x7E, 0x22)   # amber
COLOUR_LIGHT_BG   = RGBColor(0xEC, 0xF0, 0xF1)   # light grey (unused in text)
COLOUR_BODY       = RGBColor(0x2C, 0x3E, 0x50)   # near-black


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_background(cell, hex_colour: str):
    """Fill a table cell with a background colour (hex without #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = COLOUR_HEADING if level == 1 else COLOUR_SUBHEADING
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)


def add_body(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOUR_BODY
    run.font.bold = bold
    run.font.italic = italic


def add_labelled(doc: Document, label: str, value: str) -> None:
    """Bold label followed by normal value on the same paragraph."""
    p = doc.add_paragraph()
    r1 = p.add_run(f'{label}: ')
    r1.font.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = COLOUR_BODY
    r2 = p.add_run(value)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = COLOUR_BODY


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOUR_BODY


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOUR_BODY


def add_note(doc: Document, text: str, kind: str = 'tip') -> None:
    """Coloured note box using a single-cell table."""
    colour_map = {'tip': 'D5E8D4', 'warn': 'FFE6CC', 'info': 'DAE8FC'}
    fill = colour_map.get(kind, 'F5F5F5')
    label_map = {'tip': 'TIP', 'warn': 'NOTE', 'info': 'INFO'}
    label = label_map.get(kind, 'NOTE')

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    _set_cell_background(cell, fill)
    para = cell.paragraphs[0]
    r1 = para.add_run(f'{label}  ')
    r1.font.bold = True
    r1.font.size = Pt(10)
    r2 = para.add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph()  # spacer


def add_two_col_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, hdr in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_background(cell, '1A5376')
        run = cell.paragraphs[0].add_run(hdr)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        fill = 'F2F2F2' if r_idx % 2 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            _set_cell_background(cell, fill)
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.size = Pt(10)
            run.font.color.rgb = COLOUR_BODY

    doc.add_paragraph()  # spacer after table


def add_demo_section(
    doc: Document,
    title: str,
    url: str,
    objective: str,
    navigation: str,
    steps: list[str],
    talking_points: list[str],
    notes: list[tuple[str, str]] | None = None,
) -> None:
    """Render a complete demo section."""
    add_heading(doc, title, level=2)
    add_labelled(doc, 'URL', url)
    add_labelled(doc, 'Navigation', navigation)
    doc.add_paragraph()

    add_heading(doc, 'Objective', level=3)
    add_body(doc, objective)

    add_heading(doc, 'Demonstration Steps', level=3)
    for step in steps:
        add_numbered(doc, step)

    add_heading(doc, 'Talking Points', level=3)
    for point in talking_points:
        add_bullet(doc, point)

    if notes:
        for kind, text in notes:
            add_note(doc, text, kind)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------

def build_cover(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Transmission Module')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOUR_HEADING

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Demonstration Script')
    r.font.size = Pt(18)
    r.font.color.rgb = COLOUR_SUBHEADING

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = line.add_run('Sustainable Energy Now  ·  Siren Web Platform')
    r2.font.size = Pt(11)
    r2.font.color.rgb = COLOUR_BODY

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = date_para.add_run(date.today().strftime('%B %Y'))
    r3.font.size = Pt(11)
    r3.font.color.rgb = COLOUR_BODY

    doc.add_page_break()


def build_introduction(doc: Document) -> None:
    add_heading(doc, 'Introduction', level=1)
    add_body(doc, (
        'This document provides a structured walkthrough for demonstrating the Transmission module '
        'of the Siren Web platform. It covers each major view in a logical sequence, from the '
        'high-level dashboard and map views through to detailed infrastructure management and the '
        'Clean Energy Link (CEL) programme tools.'
    ))
    doc.add_paragraph()
    add_body(doc, (
        'Each section of the script includes: the URL and navigation path to reach the view; the '
        'objective of the demonstration; numbered steps to follow on screen; and key talking points '
        'for the presenter.'
    ))
    doc.add_paragraph()

    add_heading(doc, 'Module Overview', level=2)
    add_body(doc, (
        'The Transmission module manages the infrastructure that connects generation facilities to '
        'the South West Interconnected System (SWIS) grid. Key concepts:'
    ))
    add_bullet(doc, 'Terminals — substations where voltage is transformed and lines or facilities connect.')
    add_bullet(doc, 'Grid Lines — physical transmission and distribution lines between terminals.')
    add_bullet(doc, 'CEL Programmes — planned transmission investment programmes that unlock capacity for new generation.')
    add_bullet(doc, 'Viability — a scored assessment of how well a facility is positioned to benefit from CEL infrastructure.')
    doc.add_paragraph()

    add_heading(doc, 'Demonstration Sequence', level=2)
    add_two_col_table(doc,
        headers=['#', 'View', 'Purpose', 'Approx. Time'],
        rows=[
            ['1', 'Terminals Dashboard', 'System health overview', '3 min'],
            ['2', 'Grid Map', 'Geographic layout of network', '4 min'],
            ['3', 'All Terminals', 'Terminal list and detail', '4 min'],
            ['4', 'All Grid Lines', 'Grid line inventory and filters', '3 min'],
            ['5', 'Health Check', 'Data integrity diagnostics', '2 min'],
            ['6', 'CEL Transmission Map', 'CEL infrastructure on map', '4 min'],
            ['7', 'CEL Programs', 'Programme and stage management', '4 min'],
            ['8', 'Infrastructure Network', 'Full topology diagram', '3 min'],
        ]
    )

    add_heading(doc, 'Before You Start', level=2)
    add_bullet(doc, 'Log in to the Siren Web platform with an authenticated account.')
    add_bullet(doc, 'Ensure sample terminals, grid lines, and at least one CEL programme are loaded in the database.')
    add_bullet(doc, 'Set the scenario and demand year to a representative planning year using the session selector.')
    add_bullet(doc, 'Open the Transmission menu in the top navigation bar — all views are accessible from there.')
    add_note(doc, (
        'The Grid Map and CEL Transmission Map require facilities and terminals to have latitude/longitude '
        'coordinates stored in the database. Verify this before the demonstration.'
    ), kind='warn')
    doc.add_page_break()


def build_terminals_dashboard(doc: Document) -> None:
    add_demo_section(
        doc,
        title='1  ·  Terminals Network Dashboard',
        url='/terminals/dashboard/',
        navigation='Transmission → Dashboard',
        objective=(
            'Show the audience a high-level, at-a-glance view of the transmission network health. '
            'Demonstrate how the system surfaces capacity and utilisation issues automatically, '
            'without requiring the user to inspect individual records.'
        ),
        steps=[
            'Navigate to Transmission → Dashboard.',
            'Point to the four summary cards at the top: Total Terminals, Total Capacity (MVA), '
            'Connected Grid Lines, and High Utilisation.',
            'If any High Utilisation alerts are visible, click the alert link to show the affected terminals.',
            'Scroll to the "Top Terminals by Capacity" table. Explain the columns: terminal type, '
            'primary/secondary voltage, transformer capacity, and connection counts.',
            'Scroll to the "Top Terminals by Utilisation" table. Point out any amber-highlighted rows.',
            'Describe how the Technology Breakdown card at the bottom aggregates generation types '
            'connected across all terminals.',
            'Click the "Health Check" button in the top-right to transition to the next section.',
        ],
        talking_points=[
            'The dashboard gives operators an immediate view of network stress points without '
            'navigating individual terminal records.',
            'The High Utilisation threshold is 80% of transformer capacity — anything above this '
            'signals a potential bottleneck for new connections.',
            'Unconnected terminals are flagged separately because they represent data completeness '
            'issues, not just capacity issues.',
            'Utilisation is calculated from the sum of facility capacities connected through the '
            'terminal relative to its transformer MVA rating.',
            'The total MVA and average voltage figures give a quick sense of the network\'s scale.',
        ],
        notes=[
            ('tip', 'If no high-utilisation alerts are showing, you can describe the alert behaviour '
                    'verbally: "When a terminal exceeds 80%, it appears here with a link to its detail page."'),
        ]
    )


def build_grid_map(doc: Document) -> None:
    add_demo_section(
        doc,
        title='2  ·  Grid Map',
        url='/map/',
        navigation='Transmission → Grid Map  (or top navigation → Grid Map)',
        objective=(
            'Demonstrate the interactive geographic view of the transmission network. Show how '
            'terminals and facilities are positioned on the map, how connections are visualised, '
            'and how operators can update the geographic layout.'
        ),
        steps=[
            'Navigate to Grid Map via the top navigation bar.',
            'Allow the map to load. Point out the circles (terminals/substations) and the '
            'facility icons (generation assets).',
            'Click a terminal circle to open its information panel. Show the connected facilities '
            'and grid lines listed in the panel.',
            'Click a facility icon to open its information panel. Show the terminal it connects through.',
            'Demonstrate the scenario selector at the top of the page — change the demand year '
            'and explain that this filters which facilities are shown based on their planning context.',
            'Show how to reposition a terminal or facility by dragging it. Explain that this '
            'updates the geographic coordinates stored in the database.',
            'Point out the relationship between map positions and the Infrastructure Network '
            'diagram (which uses the same underlying data).',
        ],
        talking_points=[
            'The Grid Map is the primary spatial view of the network — it answers the question '
            '"where are things connected and how?"',
            'Circles represent substations (terminals); icons represent generation facilities. '
            'Lines between them are grid connections.',
            'The scenario selector is important: different scenarios include different subsets '
            'of planned and probable facilities.',
            'Drag-and-drop repositioning updates the database directly, so the map stays '
            'accurate as the physical network evolves.',
            'This view is most useful for onboarding new team members or briefing stakeholders '
            'who need a geographic context for the pipeline.',
        ],
    )


def build_terminals_management(doc: Document) -> None:
    add_demo_section(
        doc,
        title='3  ·  Terminals Management',
        url='/terminals/',
        navigation='Transmission → All Terminals',
        objective=(
            'Walk through the full lifecycle of a terminal record: viewing the list, '
            'inspecting a terminal\'s detail, understanding its connections, '
            'and demonstrating how to add or edit a terminal.'
        ),
        steps=[
            'Navigate to Transmission → All Terminals.',
            'Show the list view. Use the search box to filter by name and demonstrate '
            'how quickly a specific substation can be located.',
            'Click a terminal name to open its detail page.',
            'On the detail page, point to the attribute section: terminal type, '
            'primary voltage, secondary voltage, transformer capacity (MVA), and owner.',
            'Scroll to the Connected Facilities section. Explain that these are generation '
            'assets whose grid connections pass through this terminal.',
            'Scroll to the Connected Grid Lines section. Show the lines that have this '
            'terminal as a "from" or "to" endpoint.',
            'Click "Back" and then use the "Add New Terminal" button (or Transmission → Add Terminal) '
            'to show the creation form. Walk through the key fields without saving.',
            'Return to a terminal detail page and show the Edit and Delete buttons.',
        ],
        talking_points=[
            'Terminals are the nodes in the network graph — every grid line has a terminal '
            'at each end, and every facility connects through a terminal.',
            'The transformer capacity (MVA) is the key figure for utilisation calculations. '
            'If this is left blank, the dashboard cannot calculate utilisation for that terminal.',
            'Voltage levels (primary/secondary) distinguish terminal types: transmission '
            'substations typically operate at 132–330 kV; zone substations at 33–132 kV.',
            'Deleting a terminal removes its endpoint associations from grid lines, so '
            'check connections before deleting.',
            'The Health Check (shown next) will surface any terminals missing capacity data '
            'or grid line connections.',
        ],
        notes=[
            ('warn', 'Demonstrate the "Connect Facility" flow from the terminal detail page if time permits — '
                     'this is the mechanism for explicitly linking a facility to a substation.'),
        ]
    )


def build_health_check(doc: Document) -> None:
    add_demo_section(
        doc,
        title='4  ·  Health Check',
        url='/terminals/health-check/',
        navigation='Transmission → Health Check',
        objective=(
            'Show how the Health Check page identifies data quality and connectivity '
            'problems across the network, enabling operators to find and fix gaps '
            'before they affect analysis or visualisations.'
        ),
        steps=[
            'Navigate to Transmission → Health Check.',
            'Walk through the categories of issues the page checks: terminals with no grid lines, '
            'terminals with no facilities, grid lines with missing terminal endpoints, and '
            'terminals with missing or low capacity data.',
            'For each issue category, click an item to navigate to the relevant record and '
            'demonstrate how it would be fixed.',
            'Explain how the Health Check complements the Dashboard: the Dashboard shows '
            'operational alerts; the Health Check shows data integrity problems.',
        ],
        talking_points=[
            'The Health Check is a diagnostic tool, not a live monitoring view. Run it '
            'after bulk data imports or when adding new infrastructure.',
            'Terminals with no connected grid lines will show zero utilisation on the dashboard — '
            'they are not contributing to the network model.',
            'Grid lines with missing terminal endpoints cannot be used in topology analysis '
            'or the Infrastructure Network diagram.',
            'Regular use of the Health Check ensures the map and network views stay accurate '
            'as the pipeline evolves.',
        ],
    )


def build_grid_lines(doc: Document) -> None:
    add_demo_section(
        doc,
        title='5  ·  Grid Lines Management',
        url='/gridlines/',
        navigation='Transmission → All Grid Lines',
        objective=(
            'Demonstrate the grid line inventory, its filter and search capabilities, '
            'and the key attributes of a grid line record.'
        ),
        steps=[
            'Navigate to Transmission → All Grid Lines.',
            'Show the list with its default view. Point out the columns: name, code, '
            'line type, voltage, from terminal, to terminal, and active status.',
            'Use the Line Type filter to show only Transmission lines (≥66 kV). '
            'Then switch to Sub-transmission. Explain the three categories.',
            'Use the voltage range filter to show lines above 132 kV.',
            'Use the Connected filter to show lines that do not yet have terminal endpoints assigned.',
            'Click a grid line name to open its detail page.',
            'On the detail page, point out the thermal capacity (MVA), length (km), owner, '
            'and from/to terminal endpoints.',
            'Click "Add New Grid Line" (or Transmission → Add Grid Line) and walk through '
            'the form fields: name, type, voltage, thermal capacity, and the from/to terminal selectors.',
        ],
        talking_points=[
            'Grid lines are the edges in the network graph. Without terminal endpoints, '
            'a grid line is recorded in the inventory but plays no role in topology analysis.',
            'Line type matters for filtering and reporting: Transmission (≥66 kV) carries '
            'bulk power between major substations; Sub-transmission (11–66 kV) distributes '
            'regionally; Distribution (<11 kV) serves end consumers.',
            'Thermal capacity (MVA) is the line\'s maximum sustained current rating — '
            'this is distinct from the transformer capacity at terminals.',
            'The Active flag lets you mark decommissioned lines as inactive without deleting '
            'the historical record.',
            'Inactive lines are excluded from network calculations but remain searchable.',
        ],
    )


def build_cel_map(doc: Document) -> None:
    add_demo_section(
        doc,
        title='6  ·  CEL Transmission Map',
        url='/cel-map/',
        navigation='Transmission → Transmission Map',
        objective=(
            'Show how the CEL Transmission Map overlays planned transmission investment '
            'programmes onto the facility map, and how operators can assess which facilities '
            'are well-positioned to benefit from future infrastructure upgrades.'
        ),
        steps=[
            'Navigate to Transmission → Transmission Map.',
            'Allow the map to load with the default Development Status colour mode.',
            'Point to the facility colour legend: grey (Proposed), blue (Planned), '
            'orange (Under Construction), green (Commissioned).',
            'Switch to CEL Viability colour mode using the radio button in the left panel.',
            'Explain that the viability tier colours show how well each facility is positioned '
            'to benefit from CEL infrastructure: High (green), Medium (amber), Low (red).',
            'Use the viability tier checkboxes to filter to High-tier facilities only. '
            'Explain the scoring criteria: proximity to CEL grid lines/terminals, '
            'capacity headroom, and development probability.',
            'Switch back to Development Status mode. Use the status checkboxes to show only '
            'Under Construction and Commissioned facilities.',
            'Click a facility marker to show its popup with name, capacity, status, and '
            'CEL viability tier.',
            'Demonstrate the scenario selector to show how the view changes for different '
            'planning years.',
        ],
        talking_points=[
            'The CEL Transmission Map answers a key planning question: which generation '
            'projects are well-placed relative to planned transmission upgrades?',
            'CEL viability is pre-computed by the CEL Viability Service — scores are stored '
            'in the database and can be recalculated when programme data changes.',
            'High-viability facilities are close to CEL infrastructure with sufficient capacity '
            'headroom. Low-viability facilities may face connection constraints even after '
            'the programme is delivered.',
            'The two colour modes serve different audiences: Development Status is for '
            'pipeline tracking; CEL Viability is for transmission planning and investment analysis.',
            'Use this view alongside the CEL Programs page to see both the geographic '
            'distribution and the programme details.',
        ],
        notes=[
            ('info', 'The CEL Transmission Map links to the Grid Map and the Infrastructure Network '
                     'view via buttons in the top-right corner — use these to navigate between '
                     'spatial views during the demonstration.'),
        ]
    )


def build_cel_programs(doc: Document) -> None:
    add_demo_section(
        doc,
        title='7  ·  CEL Programs',
        url='/cel/',
        navigation='Transmission → CEL Programs',
        objective=(
            'Walk through the Clean Energy Link programme management interface. Show '
            'how programmes and stages are structured, how infrastructure is associated '
            'with each stage, and how capacity outcomes are recorded.'
        ),
        steps=[
            'Navigate to Transmission → CEL Programs.',
            'Show the programme list. Point out the programme name, code, status, '
            'and number of stages for each entry.',
            'Click a programme name to open its detail page.',
            'On the programme detail page, scroll through the stages. For each stage, '
            'point out: new capacity (MW), existing capacity unlocked (MW), and funding status.',
            'Click a stage name to open its detail page.',
            'On the stage detail page, show the associated grid lines — these are the '
            'specific transmission lines that the stage will build or upgrade.',
            'Show the associated terminals — the substations this stage affects.',
            'Explain that the combination of grid lines and terminals associated with a stage '
            'is what the CEL Viability Service uses to score facility proximity.',
            'Navigate back to the programme list and click "Add CEL Program" '
            '(or Transmission → Add CEL Program) to show the creation form.',
            'From a programme detail page, show the "Add Stage" button and walk through '
            'the stage form fields without saving.',
        ],
        talking_points=[
            'CEL Programmes represent the real-world transmission investment plans that '
            'will unlock capacity for new generation projects.',
            'The programme/stage hierarchy reflects how these investments are actually '
            'funded and delivered: a programme may span years with multiple delivery stages.',
            'New Capacity (MW) is the generation capacity that can be connected as a '
            'result of this stage\'s infrastructure. Existing Capacity (MW) is the '
            'relief provided to currently constrained generators.',
            'Funding status on each stage is important for probability weighting — '
            'a stage that is funded has a higher likelihood of delivery than one that is proposed.',
            'The link between stages and specific grid lines/terminals is what connects '
            'the programme management view to the map and viability scoring tools.',
        ],
        notes=[
            ('tip', 'If time permits, show how adding a grid line to a stage immediately affects '
                    'which facilities are scored as High viability on the CEL Transmission Map '
                    '(after viability recalculation).'),
        ]
    )


def build_infrastructure_network(doc: Document) -> None:
    add_demo_section(
        doc,
        title='8  ·  Infrastructure Dependency Network',
        url='/infrastructure-network/',
        navigation='Facilities → Network View',
        objective=(
            'Show the full topology diagram of the transmission network, demonstrating '
            'how terminals, grid lines, and facilities are interconnected, and how '
            'bottlenecks are surfaced visually.'
        ),
        steps=[
            'Navigate to Facilities → Network View.',
            'Allow the network diagram to load. Explain the three node types: '
            'large circles are terminals; smaller circles are facilities; edges are connections.',
            'Point to any amber-highlighted terminal nodes — these are above 80% utilisation '
            'and represent potential bottlenecks.',
            'Hover over a terminal node to highlight its direct connections. '
            'Explain that this makes dependency chains easy to trace.',
            'Click a terminal node to navigate to its detail page.',
            'Use the Scenario selector to switch the planning context and show how '
            'the diagram updates to include or exclude facilities.',
            'Click "Download SVG" and explain that this export can be used in reports '
            'and presentations.',
            'Use "Reset View" to clear any active highlights.',
            'Point to the colour legend below the diagram.',
        ],
        talking_points=[
            'The Infrastructure Dependency Network is the most analytical view in the '
            'module — it shows the full system topology at once, not just a geographic layout.',
            'Bottleneck detection is automatic: any terminal above 80% utilisation is '
            'highlighted without the user having to drill down into individual records.',
            'Hovering over a node highlights its immediate neighbourhood, making it '
            'easy to understand which facilities depend on a given substation.',
            'The SVG download is useful for embedding network diagrams in PowerPoint '
            'presentations or technical reports.',
            'This view is complementary to the Grid Map: the map shows geographic position; '
            'the network view shows logical connectivity and dependency.',
        ],
        notes=[
            ('info', 'This view is accessed via the Facilities menu rather than the Transmission '
                     'menu because it covers both facilities and transmission infrastructure.'),
        ]
    )


def build_glossary(doc: Document) -> None:
    add_heading(doc, 'Quick Reference Glossary', level=1)
    add_body(doc, 'Key terms for use during or after the demonstration.')
    doc.add_paragraph()

    add_two_col_table(doc,
        headers=['Term', 'Definition'],
        rows=[
            ['Terminal',            'A substation or switching station where voltage is transformed and '
                                    'grid lines or facilities connect.'],
            ['Grid Line',           'A physical transmission or distribution line connecting two terminals '
                                    'or a terminal to a generation facility.'],
            ['Transformer Capacity\n(MVA)',
                                    'Maximum throughput rating of a terminal\'s transformer equipment.'],
            ['Thermal Capacity\n(MVA)',
                                    'Maximum sustained current rating of a grid line before thermal limits are reached.'],
            ['Utilisation',         'Actual load through a terminal as a percentage of its rated transformer capacity.'],
            ['Bottleneck',          'A terminal or line operating near or above capacity, potentially '
                                    'constraining new connections.'],
            ['CEL Programme',       'A planned transmission investment programme to unlock capacity for new generation.'],
            ['CEL Stage',           'A discrete delivery phase of a programme with defined infrastructure and capacity outcomes.'],
            ['New Capacity (MW)',   'Generation capacity that can be connected as a result of a CEL stage.'],
            ['Viability Tier',      'Classification (High / Medium / Low) of how well a facility is positioned '
                                    'to benefit from a CEL stage.'],
            ['Transmission',        'Grid lines operating at ≥66 kV carrying bulk power between major substations.'],
            ['Sub-transmission',    'Grid lines at 11–66 kV distributing power regionally from zone substations.'],
            ['Distribution',        '<11 kV lines serving end consumers and small generators.'],
        ]
    )


def build_faq(doc: Document) -> None:
    add_heading(doc, 'Anticipated Questions', level=1)
    add_body(doc, (
        'Common questions that arise during demonstrations, with suggested responses.'
    ))
    doc.add_paragraph()

    faqs = [
        (
            'How does the viability score actually work?',
            'The CEL Viability Service calculates a numeric score based on three inputs: '
            'distance from the facility to the nearest CEL stage grid line or terminal; '
            'available capacity headroom at the relevant terminal; and the development status '
            'probability of the facility. The score is then bucketed into High, Medium, or Low tiers. '
            'Scores are pre-computed and stored — they can be recalculated after programme data changes.'
        ),
        (
            'Can I use this with live AEMO data?',
            'The Transmission module manages infrastructure data (terminals, grid lines, CEL programmes). '
            'It is separate from the SCADA and market data feeds used by the RET Dashboard. '
            'Facility capacity and status data can be kept current by updating facility records in Powermatch.'
        ),
        (
            'What happens if I delete a terminal that has facilities connected to it?',
            'Deleting a terminal removes its endpoint associations from grid lines, but does not delete '
            'the facilities or their grid line records. The Health Check will then flag those grid lines '
            'as missing endpoints. Always check the terminal\'s connections before deleting.'
        ),
        (
            'How do I record a grid upgrade that changes a line\'s capacity?',
            'Edit the grid line record and update the Thermal Capacity (MVA) field. '
            'If the upgrade also affects a terminal\'s transformer capacity, edit the terminal record '
            'and update the Transformer Capacity (MVA) field. The dashboard utilisation figures '
            'will reflect the new capacity immediately.'
        ),
        (
            'How often should we run the Health Check?',
            'Run the Health Check after any bulk data import or when adding a significant number '
            'of new terminals or grid lines. It is also good practice to run it before generating '
            'any reports that depend on the network topology.'
        ),
        (
            'Can the network diagram handle very large networks?',
            'The Infrastructure Dependency Network diagram is rendered as an interactive SVG '
            'in the browser. For very large networks (hundreds of nodes), rendering may be slow. '
            'The scenario filter helps by limiting the facilities shown to a specific planning context.'
        ),
    ]

    for question, answer in faqs:
        p = doc.add_paragraph()
        r1 = p.add_run(f'Q: {question}')
        r1.font.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = COLOUR_HEADING
        add_body(doc, f'A: {answer}')
        doc.add_paragraph()


def build_footer_notes(doc: Document) -> None:
    add_heading(doc, 'Demonstration Preparation Checklist', level=1)
    add_body(doc, 'Complete these steps before the demonstration session.')
    doc.add_paragraph()

    items = [
        'Confirm login credentials for an authenticated account.',
        'Verify at least 5 terminals with transformer capacity values are loaded.',
        'Verify at least 3 grid lines with terminal endpoints assigned.',
        'Verify at least one CEL programme with at least two stages.',
        'Verify at least one stage has associated grid lines and terminals.',
        'Confirm at least one facility has CEL viability scores calculated.',
        'Check that terminals and facilities have latitude/longitude for map views.',
        'Set the session scenario and demand year before starting.',
        'Test the Grid Map and CEL Transmission Map load correctly.',
        'Test the Infrastructure Network diagram renders without errors.',
        'Clear any test data or draft records that should not appear on screen.',
    ]
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        run.font.color.rgb = COLOUR_BODY

    doc.add_paragraph()
    add_labelled(doc, 'Support contact', 'modelling.lead@sen.asn.au')
    add_labelled(doc, 'Full help documentation', '/help/transmission/')
    add_labelled(doc, 'Document generated', date.today().strftime('%d %B %Y'))


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    build_cover(doc)
    build_introduction(doc)
    build_terminals_dashboard(doc)
    build_grid_map(doc)
    build_terminals_management(doc)
    build_health_check(doc)
    build_grid_lines(doc)
    build_cel_map(doc)
    build_cel_programs(doc)
    build_infrastructure_network(doc)
    build_glossary(doc)
    build_faq(doc)
    build_footer_notes(doc)

    output_path = os.path.join(os.path.dirname(__file__), 'transmission_demo_script.docx')
    doc.save(output_path)
    print(f'Document saved: {output_path}')


if __name__ == '__main__':
    main()
